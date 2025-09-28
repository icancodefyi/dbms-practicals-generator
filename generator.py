from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import time

practical_sql_operators_data = {
    'title': 'Study of operators in SQL',
    'aim': (
        'To study and apply various operators in SQL and understand their role in data manipulation and querying.'
    ),
    'theory': (
        "Operators in SQL are special symbols or keywords used to perform operations on data, allowing users to carry out arithmetic, comparison, logical, and set-based tasks directly within queries.\n\n"

        "The common types of SQL operators include:\n"
        "• Arithmetic operators (e.g., +, -, *, /, %): Perform mathematical operations on numeric values within columns or expressions.\n"
        "• Comparison operators (e.g., =, >, <, >=, <=, !=, <>): Compare two values, returning TRUE or FALSE, and are essential in WHERE clause conditions.\n"
        "• Logical operators (AND, OR, NOT): Combine multiple conditions in a WHERE clause to refine result sets based on complex criteria.\n"
        "• Set operators (UNION, INTERSECT, EXCEPT): Combine results from multiple SELECT statements, supporting advanced data analysis.\n"
        "• Special operators (BETWEEN, IN, LIKE, IS NULL): Offer versatile matching, range checks, and pattern-based queries to enhance SQL expressiveness.\n\n"

        "Effective use of operators helps filter data, join tables, combine results, and execute powerful calculations directly in the database. Mastering SQL operators enables users to write efficient, readable, and precise queries that provide meaningful insights and answer complex business questions.\n\n"

        "For example, arithmetic operators can total order values, comparison and logical operators let you select records meeting multiple criteria, and set operators can aggregate data from various sources—all from within SQL statements."
    ),
    'execution': [
        {
            'step': '1.      Arithmetic Operators',
            'code': (
                "SELECT 5 + 3 AS Sum, 10 - 2 AS Difference, 4 * 3 AS Product, 20 / 5 AS Quotient, 10 % 3 AS Remainder;"
            )
        },
        {
            'step': '2.      Comparison Operators',
            'code': (
                "SELECT Title, Author, PublishedYear FROM Books WHERE PublishedYear >= 2020 AND Author != 'J. Smith';"
            )
        },
        {
            'step': '3.      Logical Operators',
            'code': (
                "SELECT * FROM Books WHERE Genre = 'Education' OR PublishedYear < 2021;"
            )
        },
        {
            'step': '4.      Set Operators',
            'code': (
                "SELECT BookID, Title FROM Books WHERE Genre = 'Reference' UNION SELECT BookID, Title FROM Books WHERE PublishedYear > 2021;"
            )
        },
        {
            'step': '5.      Special Operators',
            'code': (
                "SELECT * FROM Books WHERE Title LIKE 'A%' OR Genre IN ('Education', 'Reference') AND PublishedYear BETWEEN 2019 AND 2023;"
            )
        }
    ],
    'output_description': (
        "Each SQL command demonstrates use of a specific type of operator, showing arithmetic calculations, comparison-based selections, logical condition filtering, set union, and special operator query functions."
    ),
    'images': [
        ('images/image.png', 'Arithmetic Operator Output'),
        ('images/image.png', 'Comparison Operator Output'),
        ('images/image.png', 'Logical Operator Output'),
        ('images/image.png', 'Set Operator Output'),
        ('images/image.png', 'Special Operator Output'),
    ],
    'outcomes': [
        'Explored SQL operators for efficient query writing.',
        'Applied arithmetic, comparison, logical, set, and special operators in practical SQL statements.',
        'Interpreted outputs to understand operator function and effect in database manipulation.'
    ],
    'conclusion': (
        'This practical improved skills in constructing advanced SQL queries using operators, supporting powerful data manipulation and extraction strategies.'
    )
}

class PracticalDocGenerator:
    def __init__(self, practical_number):
        self.doc = Document()
        self.practical_number = practical_number
        self.figure_count = 0
        self.set_page_margins()
        self.set_default_style()
        self.add_page_number_footer()

    def set_page_margins(self):
        section = self.doc.sections[-1]
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    def set_default_style(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

    def add_page_number_footer(self):
        section = self.doc.sections[-1]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')  # creates a new element
        fldChar.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def add_main_heading(self, text):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(19.5)
        run.font.name = 'Times New Roman'
        run.bold = True
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run.underline = True
        para.space_after = Pt(12)

    def add_sub_heading(self, text):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        run.bold = True
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.space_after = Pt(8)

    def add_paragraph(self, text):
        # Check if the text contains bullet points
        if '•' in text or text.startswith('-'):
            self.add_formatted_text_with_bullets(text)
        else:
            para = self.doc.add_paragraph(text)
            para.style = 'Normal'
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.space_after = Pt(10)
            para.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def add_formatted_text_with_bullets(self, text):
        """Handle text with bullet points and regular paragraphs"""
        lines = text.split('\n')
        current_paragraph = []
        
        for line in lines:
            line = line.strip()
            if not line:  # Empty line
                if current_paragraph:
                    # Add the accumulated paragraph
                    para_text = '\n'.join(current_paragraph)
                    para = self.doc.add_paragraph(para_text)
                    para.style = 'Normal'
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    para.space_after = Pt(10)
                    current_paragraph = []
                continue
                
            if line.startswith('•') or line.startswith('-'):
                # First, add any accumulated regular paragraph
                if current_paragraph:
                    para_text = '\n'.join(current_paragraph)
                    para = self.doc.add_paragraph(para_text)
                    para.style = 'Normal'
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    para.space_after = Pt(6)
                    current_paragraph = []
                
                # Add bullet point
                bullet_text = line[1:].strip()  # Remove bullet character
                bullet_para = self.doc.add_paragraph(bullet_text, style='List Bullet')
                bullet_para.space_after = Pt(3)
            else:
                # Regular text line
                current_paragraph.append(line)
        
    def add_step_description(self, step_text):
        """Add a formatted step description"""
        para = self.doc.add_paragraph(step_text)
        para.style = 'Normal'
        run = para.runs[0]
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = True
        run.font.color.rgb = RGBColor(102, 51, 0)  # Brown color for steps
        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        para.space_before = Pt(10)
        para.space_after = Pt(4)

    def add_bullet_list(self, items):
        for item in items:
            para = self.doc.add_paragraph(item, style='List Bullet')
            para.space_after = Pt(6)

    def add_code_block(self, code_text):
        para = self.doc.add_paragraph()
        run = para.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(12)

        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:val'), 'clear')
        shading_elm.set(qn('w:color'), 'auto')
        shading_elm.set(qn('w:fill'), 'F2F2F2')  # light gray background
        para._element.get_or_add_pPr().append(shading_elm)

        pBorders = OxmlElement('w:pBdr')
        leftBorder = OxmlElement('w:left')
        leftBorder.set(qn('w:val'), 'single')
        leftBorder.set(qn('w:sz'), '6')
        leftBorder.set(qn('w:space'), '4')
        leftBorder.set(qn('w:color'), 'D9D9D9')
        pBorders.append(leftBorder)
        para._element.get_or_add_pPr().append(pBorders)

        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.left_indent = Cm(0.75)
        para.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def add_image_with_caption(self, image_path, caption):
        self.figure_count += 1
        self.doc.add_picture(image_path, width=Inches(5))
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        caption_text = f'Fig {self.practical_number}.{self.figure_count} {caption}'
        cap_para = self.doc.add_paragraph(caption_text)
        cap_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cap_run = cap_para.runs[0]
        cap_run.italic = True
        cap_run.font.name = 'Times New Roman'
        cap_run.font.size = Pt(12)
        cap_para.space_after = Pt(19.5)

    def save(self, filename):
        # Add a small delay to ensure any file handles are released
        time.sleep(0.2)
        
        # Try to save with error handling
        max_attempts = 3
        for attempt in range(max_attempts):
            try:
                self.doc.save(filename)
                print(f"Successfully saved: {filename}")
                break
            except PermissionError as e:
                if attempt < max_attempts - 1:
                    print(f"File is in use, waiting... (attempt {attempt + 1}/{max_attempts})")
                    print("Please close the file if it's open in Word or another application.")
                    time.sleep(2)
                else:
                    # If still can't save, try with a different name
                    timestamp = int(time.time())
                    backup_name = f"{filename.split('.')[0]}_{timestamp}.docx"
                    try:
                        self.doc.save(backup_name)
                        print(f"Saved as {backup_name} instead (original file was in use)")
                    except Exception as backup_error:
                        print(f"Error saving file: {backup_error}")
                        raise
            except Exception as e:
                print(f"Unexpected error: {e}")
                raise

    def close_open_file(self, filename):
        """Check if file exists and warn user"""
        if os.path.exists(filename):
            print(f"Warning: {filename} already exists. Please close it if open in Word or other applications.")
            print("Attempting to save in 3 seconds... Press Ctrl+C to cancel if you need to close the file.")
            time.sleep(3)

def generate_practical_doc(practical_num, data):
    doc_gen = PracticalDocGenerator(practical_num)
    doc_gen.add_main_heading(f"Practical {practical_num}: {data['title']}")

    doc_gen.add_sub_heading("Aim:")
    doc_gen.add_paragraph(data['aim'])

    doc_gen.add_sub_heading("Theory:")
    doc_gen.add_paragraph(data['theory'])

    doc_gen.add_sub_heading("Execution:")
    for step in data['execution']:
        # step is a dict with 'step' and 'code' keys
        doc_gen.add_step_description(step['step'])  # Formatted step description
        doc_gen.add_code_block(step['code'])  # Code block for sql commands

    doc_gen.add_sub_heading("Output:")
    doc_gen.add_paragraph(data['output_description'])

    for img_path, caption in data['images']:
        doc_gen.add_image_with_caption(img_path, caption)

    doc_gen.add_sub_heading("Lab Outcomes Achieved:")
    doc_gen.add_bullet_list(data['outcomes'])

    doc_gen.add_sub_heading("Conclusion:")
    doc_gen.add_paragraph(data['conclusion'])

    filename = f"Practical_{practical_num}.docx"
    doc_gen.save(filename)
    return filename

if __name__ == "__main__":
    filename = generate_practical_doc(5, practical_sql_operators_data)
    print(f"Generated document: {filename}")
